#%%
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
from docx import Document
import re, os, tiktoken, openai, yt_dlp, youtube_episode

#%%
#Define parameters
video_url = youtube_episode.url #saved in another file to prevent saving URL directly in committed script
transcript_folder = "transcripts"

max_tokens_input=8000 #determines the token size of transcription chunks to be fed into OpenAI's API at a time. Must be under input token limit.
max_tokens_output = max_tokens_input+4000

#%%
# Initialize yt-dlp without downloading the video
with yt_dlp.YoutubeDL({'quiet': True}) as ydl:
    info = ydl.extract_info(video_url, download=False)
    # print("Title:", info.get('title'))
    # print("Uploader:", info.get('uploader'))
    # print("Upload Date:", info.get('upload_date'))

def clean_name(name):
    name = re.sub(r'\s+', ' ', name).strip()
    name = name.replace('&', 'And').title()
    name = name.replace(':', ',')
    name = name.replace('-', ',')
    return name

video_title = clean_name(info.get('title'))
channel_name = clean_name(info.get('uploader'))
#%%
# Example video URL or ID
video_id = video_url.split('v=')[-1]

# Get transcript
# transcript = YouTubeTranscriptApi.get_transcript(video_id)
# cookies = {"CONSENT": "YES+1"}  # minimal consent cookie
# transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=["en"], cookies=cookies)
# Get the TranscriptList object
api = YouTubeTranscriptApi()

# 1. Fetch a transcript (defaults to English, prefers manual)
transcript = api.fetch(video_id, languages=['en'])

# 2. To inspect available transcripts
transcript_list = api.list(video_id)

# For example, choose by language
# transcript = transcript_list.find_transcript(['en']).fetch()
# print(transcript)

gen = transcript_list.find_generated_transcript(['en'])
auto_transcript = gen.fetch()

# Join into one string
full_text = " ".join(snippet.text for snippet in transcript)

print(full_text)

# full_text = ' '.join([entry['text'] for entry in transcript])
# print(full_text)
# Please edit the following text passage by adding commas, period, and punctuation as necessary.
#%%
#Fetch OpenAI API key
api_key = os.getenv("OPENAI_API_KEY") #This key needs to be created on OpenAI's page, then saved as an environmental variable on your computer




#%%
# Function to split long transcription into smaller chunks
def split_into_chunks(text, max_tokens=max_tokens_input):
    encoder = tiktoken.get_encoding("cl100k_base")  # GPT-4's tokenizer
    sentences = text.split('. ')  # Split roughly by sentence ends
    chunks = []
    current_chunk = ''
    
    for sentence in sentences:
        # Reattach the period and ensure spacing
        sentence = sentence.strip()
        if sentence and not sentence.endswith('.'):
            sentence += '.'
        
        # Check token length if we add this sentence
        prospective_chunk = current_chunk + ' ' + sentence if current_chunk else sentence
        if len(encoder.encode(prospective_chunk)) <= max_tokens:
            current_chunk = prospective_chunk
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks
#%%

# Function to clean up each chunk of transcription using GPT-4o-mini API
def clean_transcription_chunk(raw_transcription_chunk, max_tokens=max_tokens_output, prior_final_lines=None):
    prompt = f"""
        You are helping to clean and format a transcript from a YouTube video of {channel_name}. Below is a new chunk of raw transcript text. Please edit it to make it readable, with correct grammar, punctuation, speaker formatting, and paragraph breaks.
        """

    if prior_final_lines is None:
        prompt += f"""
        Here is the transcript to edit:\n\n{raw_transcription_chunk}
        """
    else:
        prompt += f"""
        This section follows directly after the previous transcript section, which ended with the following line:\n\n{prior_final_lines}

        The new section is likely to, but may not necessarily, begin with the same speaker as the last speaker above. Here's the raw transcript to edit:\n\n{raw_transcription_chunk}
        """

    client = openai.OpenAI()  # Automatically picks up your API key from environment or config

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=max_tokens,
        temperature=0.7,
        top_p=1.0,
        n=1
    )

    cleaned_text = response.choices[0].message.content.strip()

    return cleaned_text

#%%
# Function to process a long transcription (split into chunks, clean each, and combine them)
def clean_long_transcription(long_transcription, max_tokens_input=max_tokens_input,max_tokens_output=max_tokens_output):
    # Split transcription into manageable chunks
    chunks = split_into_chunks(long_transcription,max_tokens_input)
    
    cleaned_chunks=[]
    final_prior_lines=''

    # Clean up each chunk and combine results
    for i, chunk in enumerate(chunks):
        if i==0:
            cleaned_chunk = clean_transcription_chunk(chunk, max_tokens_output)
        else:
            cleaned_chunk = clean_transcription_chunk(chunk,max_tokens_output,final_prior_lines)
        
        # Extract the final line to pass to the next chunk
        lines = [line.strip() for line in cleaned_chunk.strip().split('\n') if line.strip()]
        if lines:
            # final_prior_lines = "\n".join(lines[-1:])
            final_prior_lines = lines[-1:]
        lines=[]

        cleaned_chunks.append(cleaned_chunk)

    # cleaned_chunks = [clean_transcription_chunk(chunk) for chunk in chunks]
    cleaned_transcription = "\n".join(cleaned_chunks)
    
    return cleaned_transcription


#%%
cleaned_transcription = clean_long_transcription(full_text)

# Save the cleaned transcript to a Word document
cleaned_transcript_path_name_doc = os.path.join(transcript_folder, f"{video_title}.docx")
cleaned_doc = Document()
# cleaned_doc.add_paragraph(cleaned_transcription)
paragraph = cleaned_doc.add_paragraph()

# ChatGPT highlights select text with "**".  Use regex to split text into normal and bold parts
parts = re.split(r'(\*\*.*?\*\*)', cleaned_transcription)

for part in parts:
    if part.startswith('**') and part.endswith('**'):
        # Remove the ** markers
        clean_text = part[2:-2]
        run = paragraph.add_run(clean_text)
        run.bold = True
    else:
        # Normal text
        paragraph.add_run(part)

# Save the document
cleaned_doc.save(cleaned_transcript_path_name_doc)


"""
Next steps:
- Tinker with token limit. The ideal is as large as possible, as to minimize the risk of misidentifying the speaker and inconsistencies between chunks. But small enough that hitting the token limit is never a concern.
- determine if "final line" logic needs refining. since speakers talk for a while, it may not be necessary to extract everything since the last break between speakers.
"""

# %%
