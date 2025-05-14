
#%%
import requests
import re
import os
from urllib.parse import urljoin
import requests
import whisper
from mutagen.easyid3 import EasyID3
from mutagen.mp3 import MP3
from docx import Document
from bs4 import BeautifulSoup
import openai
import tiktoken

#%%
#Define parameters
podcast_episode_url = "<PODCAST URL>" #The URL where you can find the podcast. To date, I have only tested this on the Apple Podcast pages, specific to each episode.
download_folder = "downloads"
transcript_folder = "transcripts"
podcast_description = "The Drive podcast with Peter Attia" #The information you are giving ChatGPT about your specific podcast, in the prompt.

max_tokens_input=8000 #determines the token size of transcription chunks to be fed into OpenAI's API at a time. Must be under input token limit.
max_tokens_output = max_tokens_input+4000

#%%
def get_episode_title_from_page(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # Try Open Graph <meta property="og:title">
    og_title = soup.find("meta", property="og:title")
    if og_title and og_title.get("content"):
        return og_title["content"]

    # Try <meta name="title">
    meta_title = soup.find("meta", attrs={"name": "title"})
    if meta_title and meta_title.get("content"):
        return meta_title["content"]

    # Fall back to <title> tag
    if soup.title and soup.title.string:
        return soup.title.string.strip()

    return "untitled_episode"

podcast_title = re.sub(r'[^\x00-\x7F]+', '', get_episode_title_from_page(podcast_episode_url))
# Step 2: Collapse multiple spaces into a single space
podcast_title = re.sub(r'\s+', ' ', podcast_title).strip()
#convert "&"" to "and"
podcast_title = podcast_title.replace("&", "And").title()
podcast_title = podcast_title.replace(":", ",")
podcast_title = podcast_title.replace("-", ",")

print(podcast_title)
#%%
def download_unique_mp3s_with_titles(url, download_folder=download_folder, transcript_folder=transcript_folder, podcast_title=podcast_title):
    response = requests.get(url)
    if response.status_code != 200:
        print(f"Failed to load page: {response.status_code}")
        return

    html = response.text

    # Find all .mp3 links using regex and keep only unique ones
    raw_links = re.findall(r'https?://[^\s"\']+\.mp3', html)
    unique_links = list(set(raw_links))

    if not unique_links:
        print("No .mp3 links found.")
        return

    os.makedirs(download_folder, exist_ok=True)
    os.makedirs(transcript_folder, exist_ok=True)

    safe_titles = []

    for i, mp3_url in enumerate(unique_links):
        temp_filename = os.path.join(download_folder, f"temp_{i}.mp3")
        print(f"Downloading: {mp3_url}")

        try:
            # Download the MP3
            r = requests.get(mp3_url, stream=True)
            with open(temp_filename, 'wb') as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)

            # Extract title from metadata
            audio = MP3(temp_filename, ID3=EasyID3)
            title = audio.get('title', [f'podcast_{i+1}'])[0]
            safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-")).strip()
            safe_titles.append(safe_title)
            # final_mp3_path = os.path.join(download_folder, f"{safe_title}.mp3")
            final_mp3_path = os.path.join(download_folder, f"{podcast_title}.mp3")

            # Rename file using title
            os.rename(temp_filename, final_mp3_path)
            print(f"Saved as: {final_mp3_path}")

            # Save empty transcript file as a placeholder
            # transcript_path = os.path.join(transcript_folder, f"{safe_title}.txt")
            # with open(transcript_path, 'w') as tf:
            #     tf.write(f"Transcript placeholder for: {safe_title}\n")

        except Exception as e:
            print(f"Error processing {mp3_url}: {e}")
            if os.path.exists(temp_filename):
                os.remove(temp_filename)
    return safe_titles

# Download the podcast, and save a local copy
download_unique_mp3s_with_titles(podcast_episode_url)

#%%
file_name = podcast_title
mp3_path_name = os.path.join(download_folder, f"{file_name}.mp3")
transcript_path_name_doc = os.path.join(transcript_folder, f"{file_name}.docx")

# Load Whisper model and transcribe the audio file
model = whisper.load_model("base")  # "tiny", "small", "medium", "large" also work
result = model.transcribe(mp3_path_name)

long_transcription = result['text']
print(long_transcription)

# Save the transcript to a Word document
doc = Document()
# doc.add_heading(file_name, level=1)
doc.add_paragraph(long_transcription)
doc.save(transcript_path_name_doc)

#save the file also to a text file
# transcript_path_name_txt= os.path.join(transcript_folder, f"{file_name}.txt")
# with open(transcript_path_name_txt, "w", encoding="utf-8") as f:
#     f.write(long_transcription)

#delete mp3 file
os.remove(mp3_path_name)

#%%
#Fetch OpenAI API key
api_key = os.getenv("OPENAI_API_KEY") #This key needs to be created on OpenAI's page, then saved as an environmental variable on your computer




#%%
# Function to split long transcription into smaller chunks
def split_into_chunks(text, max_tokens=max_tokens_input):
    encoder = tiktoken.get_encoding("cl100k_base")  # GPT-4's tokenizer
    sentences = text.split('. ')  # Split roughly by sentence ends
    chunks = []
    current_chunk = ''
    
    for sentence in sentences:
        # Reattach the period and ensure spacing
        sentence = sentence.strip()
        if sentence and not sentence.endswith('.'):
            sentence += '.'
        
        # Check token length if we add this sentence
        prospective_chunk = current_chunk + ' ' + sentence if current_chunk else sentence
        if len(encoder.encode(prospective_chunk)) <= max_tokens:
            current_chunk = prospective_chunk
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks
#%%

# Function to clean up each chunk of transcription using GPT-4o-mini API
def clean_transcription_chunk(raw_transcription_chunk, max_tokens=max_tokens_output, prior_final_lines=None, podcast_description=podcast_description):
    prompt = f"You are helping to clean and format a transcript from an episode of {podcast_description}. Below is a new chunk of raw transcript text. Please edit it to make it readable, with correct grammar, punctuation, speaker formatting, and paragraph breaks."

    if prior_final_lines is None:
        prompt += f"""
        Format the output as a dialogue, like this:
        Host: What is your favorite color?
        Guest: My favorite color is blue.
        Host: Cool.

        Here is the transcript to edit:\n\n{raw_transcription_chunk}
        """
    else:
        prompt += f"""
        This section follows directly after the previous transcript section, which ended with the following line:\n\n{prior_final_lines}

        The new section is likely to, but may not necessarily, begin with the same speaker as the last speaker above. Here's the raw transcript to edit:\n\n{raw_transcription_chunk}
        """

    client = openai.OpenAI()  # Automatically picks up your API key from environment or config

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=max_tokens,
        temperature=0.7,
        top_p=1.0,
        n=1
    )

    cleaned_text = response.choices[0].message.content.strip()

    return cleaned_text

#%%
# Function to process a long transcription (split into chunks, clean each, and combine them)
def clean_long_transcription(long_transcription, max_tokens_input=max_tokens_input,max_tokens_output=max_tokens_output):
    # Split transcription into manageable chunks
    chunks = split_into_chunks(long_transcription,max_tokens_input)
    
    cleaned_chunks=[]
    final_prior_lines=''

    # Clean up each chunk and combine results
    for i, chunk in enumerate(chunks):
        if i==0:
            cleaned_chunk = clean_transcription_chunk(chunk, max_tokens_output)
        else:
            cleaned_chunk = clean_transcription_chunk(chunk,max_tokens_output,final_prior_lines)
        
        # Extract the final line to pass to the next chunk
        lines = [line.strip() for line in cleaned_chunk.strip().split('\n') if line.strip()]
        if lines:
            # final_prior_lines = "\n".join(lines[-1:])
            final_prior_lines = lines[-1:]
        lines=[]

        cleaned_chunks.append(cleaned_chunk)

    # cleaned_chunks = [clean_transcription_chunk(chunk) for chunk in chunks]
    cleaned_transcription = "\n".join(cleaned_chunks)
    
    return cleaned_transcription


#%%
cleaned_transcription = clean_long_transcription(long_transcription)

# Save the cleaned transcript to a Word document
cleaned_transcript_path_name_doc = os.path.join(transcript_folder, f"{file_name} - CLEANED.docx")
cleaned_doc = Document()
cleaned_doc.add_paragraph(cleaned_transcription)
cleaned_doc.save(cleaned_transcript_path_name_doc)