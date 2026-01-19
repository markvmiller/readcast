"""Word document generation for transcripts."""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict


class DocumentGenerator:
    """Generates Word documents from transcripts."""
    
    def __init__(self, transcript_folder: str = "transcripts"):
        self.transcript_folder = transcript_folder
        os.makedirs(transcript_folder, exist_ok=True)
    
    def create_document(self, episode_title: str, podcast_title: str, speakers: Dict[str, List[str]], 
                       transcript_text: str, cleaned: bool = True) -> str:
        """
        Create a Word document from transcript.
        
        Args:
            episode_title: Title of the episode
            podcast_title: Title of the podcast
            speakers: Dictionary with speaker information
            transcript_text: The transcript content
            cleaned: Whether this is a cleaned transcript
            
        Returns:
            Path to created document
        """
        # Create filename
        suffix = " - CLEANED" if cleaned else ""
        safe_title = "".join(c for c in episode_title if c.isalnum() or c in (" ", "_", "-")).strip()
        filename = f"{safe_title}{suffix}.docx"
        filepath = os.path.join(self.transcript_folder, filename)
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(episode_title, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add podcast info
        doc.add_heading("Podcast Information", level=2)
        doc.add_paragraph(f"Podcast: {podcast_title}")
        
        # Add speaker information
        doc.add_paragraph(f"Host: {speakers['host']}")
        
        if speakers.get("cohosts"):
            cohosts_str = ", ".join(speakers["cohosts"])
            doc.add_paragraph(f"Co-hosts: {cohosts_str}")
        
        if speakers.get("guests"):
            guests_str = ", ".join(speakers["guests"])
            doc.add_paragraph(f"Guest(s): {guests_str}")
        
        # Add separator
        doc.add_paragraph("\n" + "="*50 + "\n")
        
        # Add transcript content
        doc.add_heading("Transcript", level=2)
        self._add_formatted_text(doc, transcript_text)
        
        # Save document
        doc.save(filepath)
        print(f"Document saved: {filepath}")
        
        return filepath
    
    def _add_formatted_text(self, doc: Document, text: str) -> None:
        """Add text with formatting (handles bold text marked with **)."""
        paragraph = doc.add_paragraph()
        
        # Split text by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                clean_text = part[2:-2]
                run = paragraph.add_run(clean_text)
                run.bold = True
            else:
                # Normal text
                paragraph.add_run(part)
    
    def get_document_path(self, episode_title: str, cleaned: bool = True) -> str:
        """Get the expected document path for an episode."""
        suffix = " - CLEANED" if cleaned else ""
        safe_title = "".join(c for c in episode_title if c.isalnum() or c in (" ", "_", "-")).strip()
        filename = f"{safe_title}{suffix}.docx"
        return os.path.join(self.transcript_folder, filename)
