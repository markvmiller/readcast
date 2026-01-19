#%%
import requests, re, os, json, html, whisper, openai, tiktoken
from urllib.parse import urljoin
from mutagen.easyid3 import EasyID3
from mutagen.mp3 import MP3
from docx import Document
from bs4 import BeautifulSoup, Comment
import podcast_episode 

#%%
#Define parameters
#The URL where you can find the podcast. To date, I have only tested this on the Apple Podcast pages, specific to each episode.
podcast_episode_url = podcast_episode.url #saved in another file to prevent saving URL directly in committed script
download_folder = "downloads"
transcript_folder = "transcripts"
openai_model="gpt-4o-mini"

#Fetch OpenAI API key
api_key = os.getenv("OPENAI_API_KEY") #This key needs to be created on OpenAI's page, then saved as an environmental variable on your computer
if not api_key:
    raise RuntimeError("OpenAI API key must be supplied (env var or argument).")

max_tokens_input=8000 #determines the token size of transcription chunks to be fed into OpenAI's API at a time. Must be under input token limit.
max_tokens_output = max_tokens_input+4000

#%%
def get_podcast_channel_from_page(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # Look through all JSON-LD <script> tags
    for script_tag in soup.find_all("script", type="application/ld+json"):
        try:
            data = json.loads(script_tag.string)

            # Handle single or multiple objects in the JSON
            if isinstance(data, dict):
                items = [data]
            elif isinstance(data, list):
                items = data
            else:
                continue

            for item in items:
                # Look for PodcastEpisode object with partOfSeries
                if item.get("@type") == "PodcastEpisode":
                    series = item.get("partOfSeries")
                    if isinstance(series, dict):
                        name = series.get("name")
                        if name:
                            return name
        except (json.JSONDecodeError, TypeError):
            continue

    return "untitled_podcast"

podcast_title = get_podcast_channel_from_page(podcast_episode_url)
print(podcast_title)


#%%
def get_full_episode_description(url: str) -> str:
    """
    Return the complete Apple-Podcasts episode description, with paragraphs
    separated by blank lines.  Falls back gracefully if one layer is missing.
    """
    ## 1. download & decode --------------------------------------------------
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    resp.encoding = "utf-8"                     # stop ‘isnât’ mojibake
    soup = BeautifulSoup(resp.text, "html.parser")

    ## 2. easiest path – JSON-LD --------------------------------------------
    json_fallback = None
    for tag in soup.find_all("script", type="application/ld+json"):
        try:
            data = json.loads(tag.string)
        except (TypeError, json.JSONDecodeError):
            continue

        for item in (data if isinstance(data, list) else [data]):
            if item.get("@type") == "PodcastEpisode":
                desc = item.get("description")
                if desc and len(desc) > 400:   # long enough? take it.
                    return desc.strip()
                json_fallback = desc or json_fallback

    ## 3. hard path – HTML comment block ------------------------------------
    wrapper = soup.find("div", attrs={"data-testid": "paragraphs"})
    if wrapper:
        paragraphs = []

        for node in wrapper.descendants:
            # a) visible <p> tags
            if getattr(node, "name", None) == "p":
                paragraphs.append(node.get_text(" ", strip=True))

            # b) hidden HTML in comments
            elif isinstance(node, Comment):
                sub = BeautifulSoup(node, "html.parser")
                paragraphs.extend(
                    p.get_text(" ", strip=True) for p in sub.find_all("p")
                )

        # clean up & join
        text = "\n\n".join(html.unescape(p) for p in paragraphs if p)
        if text:
            return text

    ## 4. last resort – whatever JSON-LD held --------------------------------
    if json_fallback:
        return json_fallback.strip()

    return "no_description_found"

episode_description = get_full_episode_description(podcast_episode_url)
print(episode_description)
#%%
def get_episode_title_from_page(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # Try Open Graph <meta property="og:title">
    og_title = soup.find("meta", property="og:title")
    if og_title and og_title.get("content"):
        return og_title["content"]

    # Try <meta name="title">
    meta_title = soup.find("meta", attrs={"name": "title"})
    if meta_title and meta_title.get("content"):
        return meta_title["content"]

    # Fall back to <title> tag
    if soup.title and soup.title.string:
        return soup.title.string.strip()

    return "untitled_episode"

episode_title = re.sub(r'[^\x00-\x7F]+', '', get_episode_title_from_page(podcast_episode_url))
# Step 2: Collapse multiple spaces into a single space
episode_title = re.sub(r'\s+', ' ', episode_title).strip()
#convert "&"" to "and"
episode_title = episode_title.replace("&", "And").title()
episode_title = episode_title.replace(":", ",")
episode_title = episode_title.replace("-", ",")
episode_title = episode_title.replace(" ,", ",")

print(episode_title)

#%%
def download_unique_mp3s_with_titles(url, download_folder=download_folder, transcript_folder=transcript_folder, episode_title=episode_title):
    response = requests.get(url)
    if response.status_code != 200:
        print(f"Failed to load page: {response.status_code}")
        return

    html = response.text

    # Find all .mp3 links using regex and keep only unique ones
    raw_links = re.findall(r'https?://[^\s"\']+\.mp3', html)
    unique_links = list(set(raw_links))

    if not unique_links:
        print("No .mp3 links found.")
        return

    os.makedirs(download_folder, exist_ok=True)
    os.makedirs(transcript_folder, exist_ok=True)

    safe_titles = []

    for i, mp3_url in enumerate(unique_links):
        temp_filename = os.path.join(download_folder, f"temp_{i}.mp3")
        print(f"Downloading: {mp3_url}")

        try:
            # Download the MP3
            r = requests.get(mp3_url, stream=True)
            with open(temp_filename, 'wb') as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)

            # Extract title from metadata
            audio = MP3(temp_filename, ID3=EasyID3)
            title = audio.get('title', [f'podcast_{i+1}'])[0]
            safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-")).strip()
            safe_titles.append(safe_title)
            # final_mp3_path = os.path.join(download_folder, f"{safe_title}.mp3")
            final_mp3_path = os.path.join(download_folder, f"{episode_title}.mp3")

            # Rename file using title
            os.rename(temp_filename, final_mp3_path)
            print(f"Saved as: {final_mp3_path}")

            # Save empty transcript file as a placeholder
            # transcript_path = os.path.join(transcript_folder, f"{safe_title}.txt")
            # with open(transcript_path, 'w') as tf:
            #     tf.write(f"Transcript placeholder for: {safe_title}\n")

        except Exception as e:
            print(f"Error processing {mp3_url}: {e}")
            if os.path.exists(temp_filename):
                os.remove(temp_filename)
    return safe_titles

# Download the podcast, and save a local copy
download_unique_mp3s_with_titles(podcast_episode_url)

#%%
file_name = episode_title
mp3_path_name = os.path.join(download_folder, f"{file_name}.mp3")
transcript_path_name_doc = os.path.join(transcript_folder, f"{file_name}.docx")

# Load Whisper model and transcribe the audio file
model = whisper.load_model("base")  # "tiny", "small", "medium", "large" also work
result = model.transcribe(mp3_path_name)

long_transcription = result['text']
print(long_transcription)

# Save the transcript to a Word document
doc = Document()
# doc.add_heading(file_name, level=1)
doc.add_paragraph(long_transcription)
doc.save(transcript_path_name_doc)

#save the file also to a text file
# transcript_path_name_txt= os.path.join(transcript_folder, f"{file_name}.txt")
# with open(transcript_path_name_txt, "w", encoding="utf-8") as f:
#     f.write(long_transcription)

#delete mp3 file
os.remove(mp3_path_name)

#%%




#%%
# Split long transcription into smaller chunks
def split_into_chunks(text, max_tokens=max_tokens_input):
    encoder = tiktoken.get_encoding("cl100k_base")  # GPT-4's tokenizer
    sentences = text.split('. ')  # Split roughly by sentence ends
    chunks = []
    current_chunk = ''
    
    for sentence in sentences:
        # Reattach the period and ensure spacing
        sentence = sentence.strip()
        if sentence and not sentence.endswith('.'):
            sentence += '.'
        
        # Check token length if we add this sentence
        prospective_chunk = current_chunk + ' ' + sentence if current_chunk else sentence
        if len(encoder.encode(prospective_chunk)) <= max_tokens:
            current_chunk = prospective_chunk
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

#%%
#infer speaker(s) from podcast and episode titles and description:
def extract_speakers(description=episode_description, podcast_title=podcast_title, episode_title=episode_title, openai_model=openai_model, temperature=0.0):
    """
    Extracts host, optional cohosts, and guests from a podcast episode description.

    Returns a dict like:
    {
        "host": "Bob Smith",
        "cohosts": ["Jane Doe"],
        "guests": ["Alice Jones"]
    }

    Raises:
        RuntimeError if host is not found.
    """
    system_prompt = (
        "You are an assistant that extracts speakers from podcast metadata.\n"
        "Return a JSON object with these keys only:\n"
        "• host (REQUIRED, string)\n"
        "• cohosts (OPTIONAL, list of strings)\n"
        "• guests (OPTIONAL, list of strings)\n"
        "Only respond with valid JSON. Do not include extra commentary."
    )

    user_prompt = f"""
Podcast Title: {podcast_title}
Episode Title: {episode_title}
Episode Description:
{description}
    """.strip()

    client = openai.OpenAI()  # Uses API key from environment

    response = client.chat.completions.create(
        model=openai_model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=temperature
    )

    content = response.choices[0].message.content.strip()

    try:
        result = json.loads(content)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"The model's response was not valid JSON: {result}") from e

    if "host" not in result or not result["host"].strip():
        raise RuntimeError(f"No host identified in the model's output: {result}")

    return result

podcast_episode_speakers = extract_speakers()
print(podcast_episode_speakers)

#%%
podcast_host = podcast_episode_speakers["host"]
print(f"Podcast host: {podcast_host}")

podcast_description = f"{podcast_title} podcast with {podcast_host}" 
if not podcast_description.startswith("The "):
    podcast_description = "The " + podcast_description

def join_list_elements(item_list):
    """
    Turn a list of elements into a human-readable string.

    • 1 name  → "Alice"
    • 2 names → "Alice and Bob"
    • 3+      → "Alice, Bob, and Carol"
    """
    if len(item_list) == 1:
        return item_list[0]
    elif len(item_list) == 2:
        return " and ".join(item_list)
    else:
        return ", ".join(item_list[:-1]) + ", and " + item_list[-1]
    
if podcast_episode_speakers["cohosts"]:
    podcast_cohost_list = podcast_episode_speakers["cohosts"]
    podcast_cohosts = join_list_elements(podcast_cohost_list)
    print(f"Podcast cohosts: {podcast_cohosts}")

    podcast_description+= " and " if len(podcast_cohost_list)==1 else ", "
    podcast_description+= f"{podcast_cohosts}"

if podcast_episode_speakers["guests"]:
    podcast_guest_list = podcast_episode_speakers["guests"]
    podcast_guests = join_list_elements(podcast_guest_list)
    print(f"Podcast guest(s): {podcast_guests}")

    podcast_description += f", featuring guest"
    podcast_description+= " " if len(podcast_guest_list)==1 else "s "
    podcast_description+= f"{podcast_guests}"

print(podcast_description)

#%%
# Function to clean up each chunk of transcription using GPT-4o-mini API
def clean_transcription_chunk(raw_transcription_chunk, max_tokens=max_tokens_output, prior_final_lines=None, podcast_description=podcast_description, openai_model=openai_model):
    prompt = f"""
        You are helping to clean and format a transcript from an episode of {podcast_description}. Below is a new chunk of raw transcript text. Please edit it to make it readable, with correct grammar, punctuation, speaker formatting, and paragraph breaks.
        """

    if 'podcast_guests' in globals():    
        prompt +=f"""

        Format the output as a dialogue, like this:
        {podcast_host}: What is your favorite color?
        {podcast_guest_list[0]}: My favorite color is blue.
        {podcast_host}: Cool.
        
        """

    if prior_final_lines is None:
        prompt += f"""
        Here is the transcript to edit:\n\n{raw_transcription_chunk}
        """
    else:
        prompt += f"""
        This section follows directly after the previous transcript section, which ended with the following line:\n\n{prior_final_lines}

        The new section is likely to, but may not necessarily, begin with the same speaker as the last speaker above. Here's the raw transcript to edit:\n\n{raw_transcription_chunk}
        """

    client = openai.OpenAI()  # Automatically picks up your API key from environment or config

    response = client.chat.completions.create(
        model=openai_model,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=max_tokens,
        temperature=0.7,
        top_p=1.0,
        n=1
    )

    cleaned_text = response.choices[0].message.content.strip()

    return cleaned_text

#%%
# Function to process a long transcription (split into chunks, clean each, and combine them)
def clean_long_transcription(long_transcription, max_tokens_input=max_tokens_input,max_tokens_output=max_tokens_output):
    # Split transcription into manageable chunks
    chunks = split_into_chunks(long_transcription,max_tokens_input)
    
    cleaned_chunks=[]
    final_prior_lines=''

    # Clean up each chunk and combine results
    for i, chunk in enumerate(chunks):
        if i==0:
            cleaned_chunk = clean_transcription_chunk(chunk, max_tokens_output)
        else:
            cleaned_chunk = clean_transcription_chunk(chunk,max_tokens_output,final_prior_lines)
        
        # Extract the final line to pass to the next chunk
        lines = [line.strip() for line in cleaned_chunk.strip().split('\n') if line.strip()]
        if lines:
            # final_prior_lines = "\n".join(lines[-1:])
            final_prior_lines = lines[-1:]
        lines=[]

        cleaned_chunks.append(cleaned_chunk)

    # cleaned_chunks = [clean_transcription_chunk(chunk) for chunk in chunks]
    cleaned_transcription = "\n".join(cleaned_chunks)
    
    return cleaned_transcription

#%%
cleaned_transcription = clean_long_transcription(long_transcription)

# Save the cleaned transcript to a Word document
# cleaned_transcript_path_name_doc = os.path.join(transcript_folder, f"{file_name} - CLEANED.docx")
# cleaned_doc = Document()
# cleaned_doc.add_paragraph(cleaned_transcription)
# cleaned_doc.save(cleaned_transcript_path_name_doc)

#%%
# Save the cleaned transcript to a Word document
cleaned_transcript_path_name_doc = os.path.join(transcript_folder, f"{file_name} - CLEANED.docx")
cleaned_doc = Document()
# cleaned_doc.add_paragraph(cleaned_transcription)
paragraph = cleaned_doc.add_paragraph()

# ChatGPT highlights select text with "**".  Use regex to split text into normal and bold parts
parts = re.split(r'(\*\*.*?\*\*)', cleaned_transcription)

for part in parts:
    if part.startswith('**') and part.endswith('**'):
        # Remove the ** markers
        clean_text = part[2:-2]
        run = paragraph.add_run(clean_text)
        run.bold = True
    else:
        # Normal text
        paragraph.add_run(part)

# Save the document
cleaned_doc.save(cleaned_transcript_path_name_doc)
"""
Next steps:
- Clean code, and consolidate where it makes sense to do so.
- Tinker with token limit. The ideal is as large as possible, as to minimize the risk of misidentifying the speaker and inconsistencies between chunks. But small enough that hitting the token limit is never a concern.
- determine if "final line" logic needs refining. since speakers talk for a while, it may not be necessary to extract everything since the last break between speakers.
- modify prompt to look for and omit commercials
- the output doesn't always include speaker refereces. refine prompting to ensure it is included.
- Explore converting to a web UI which takes a URL as input and outputs a word file for downloading.
"""
# %%
